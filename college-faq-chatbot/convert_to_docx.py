"""Convert the BVRIT knowledge base .md to .docx format."""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
with open('../BVRIT HYDERABAD Chatbot/BVRIT_Hyderabad_Knowledge_Base.md', 'r', encoding='utf-8') as f:
    content = f.read()

for line in content.split('\n'):
    stripped = line.strip()
    if line.startswith('# '):
        doc.add_heading(stripped[2:], level=1)
    elif line.startswith('## '):
        doc.add_heading(stripped[3:], level=2)
    elif line.startswith('### '):
        doc.add_heading(stripped[4:], level=3)
    elif line.startswith('#### '):
        doc.add_heading(stripped[5:], level=4)
    elif line.startswith('---'):
        p = doc.add_paragraph()
        run = p.add_run('_' * 50)
    elif line.startswith('- '):
        doc.add_paragraph(stripped[2:], style='List Bullet')
    elif line.startswith('> '):
        doc.add_paragraph(stripped[2:], style='Intense Quote')
    elif line.startswith('```'):
        continue
    elif stripped:
        doc.add_paragraph(stripped)

doc.save('data/college_knowledge.docx')
print('DOCX created successfully at data/college_knowledge.docx')